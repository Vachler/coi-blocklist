import requests
import io
import re
from docx import Document

def get_coi_data():
    url = "https://www.coi.gov.cz/pro-spotrebitele/rizikove-e-shopy/"
    domains = set()
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        r = requests.get(url, headers=headers, timeout=30)
        found = re.findall(r'<td>([a-z0-9.-]+\.[a-z]{2,10})</td>', r.text.lower())
        for d in found:
            if d and '.' in d: domains.add(d.strip())
    except: pass
    return domains

def get_soi_data():
    url = "https://www.soi.sk/files/documents/rizikove-eshopy/zoznam-rizikovych-eshopov.docx"
    domains = set()
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        r = requests.get(url, headers=headers, timeout=30)
        if r.status_code == 200:
            doc = Document(io.BytesIO(r.content))
            full_text = ""
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
            for para in doc.paragraphs:
                full_text += para.text + " "
            found = re.findall(r'([a-z0-9.-]+\.[a-z]{2,10})', full_text.lower())
            for d in found:
                if '.' in d and not any(x in d for x in ['soi.sk', 'gov.sk', 'microsoft', 'w3.org']):
                    domains.add(d.strip().strip('.'))
    except: pass
    return domains

if __name__ == "__main__":
    all_domains = get_coi_data() | get_soi_data()
    final_list = sorted([f"||{d}^" for d in all_domains if len(d) > 3])
    if final_list:
        with open("blocklist.txt", "w", encoding="utf-8") as f:
            f.write("! Title: Oficialni CZ & SK Blocklist\n")
            f.write(f"! Total items: {len(final_list)}\n\n")
            f.write("\n".join(final_list))
